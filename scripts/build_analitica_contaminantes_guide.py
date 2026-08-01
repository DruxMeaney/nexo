from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Raiz del repositorio, deducida de la ubicacion de este script: el guion
# debe correr en cualquier clon, no solo en la maquina del autor.
ROOT = Path(__file__).resolve().parent.parent
SLIDE_DIR = ROOT / "outputs/manual-analitica-guide/presentations/analitica-contaminantes/template-inspect/source-slides"
DOCS_DIR = ROOT / "docs"
GUIDES_DIR = ROOT / "outputs/guias"
MD_OUT = DOCS_DIR / "guia_detallada_analitica_contaminantes.md"
DOCX_OUT = GUIDES_DIR / "guia_detallada_analitica_contaminantes.docx"


TITLE = "Guia detallada para presentar analitica_contaminantes.pptx"
SUBTITLE = "Contaminantes en agua y enfermedades neurodegenerativas: lectura metodologica, matematica e interpretativa de cada diapositiva"


overview_sections = [
    (
        "Tesis central de la presentacion",
        [
            "Esta presentacion no pretende cerrar conclusiones causales sobre contaminantes y enfermedades neurodegenerativas. Su objetivo es mostrar un sistema reproducible para convertir articulos cientificos en evidencia estructurada: menciones, contexto, seccion del articulo, relaciones contaminante-enfermedad y graficos de priorizacion.",
            "La idea fuerte que conviene repetir es: el conteo de palabras es solo la primera capa. La parte importante del algoritmo es que cada conteo queda unido a una ventana textual, a la seccion donde aparece y a reglas auditables que reducen falsos positivos.",
            "La presentacion avanza en tres niveles: primero explica como se construye el algoritmo; despues muestra la base de evidencia en conteos y secciones; finalmente usa matrices, burbujas, redes y clustering para priorizar la revision manual.",
        ],
    ),
    (
        "Mensaje de 30 segundos para abrir",
        [
            "Analizamos 79 articulos con un pipeline en Python que detecta contaminantes y enfermedades neurodegenerativas, pero no se queda en contar palabras. Cada mencion se contextualiza por seccion, ventana textual y pistas linguisticas. Esto permite separar una mencion bibliografica de una posible exposicion trabajada en el estudio. Las figuras no son conclusiones clinicas; son mapas de priorizacion para saber que pares contaminante-enfermedad merecen auditoria manual primero.",
        ],
    ),
    (
        "Diferencia clave que debes dominar",
        [
            "Mencion no significa evidencia. Una mencion es que una palabra o sinonimo aparecio en el texto.",
            "Entidad relevante no significa causalidad. Significa que el contaminante o la enfermedad aparecen con suficiente informacion contextual para ser registrados.",
            "Asociacion no significa mecanismo causal. En este pipeline, asociacion quiere decir que contaminante y enfermedad aparecen con proximidad textual y soporte contextual. La fuerza depende de seccion, confianza, pistas de exposicion, pistas de relacion y presencia de evidencia textual.",
            "Grafico no significa conclusion. Los graficos organizan la evidencia y ayudan a priorizar que revisar manualmente.",
        ],
    ),
    (
        "Numeros globales que conviene memorizar",
        [
            "79 articulos procesados.",
            "78 articulos con texto extraible y 1 articulo sin texto util extraible.",
            "14,739 menciones detectadas en total: 9,831 menciones de contaminantes y 4,908 menciones de enfermedades.",
            "936 relaciones contaminante-enfermedad con evidencia textual.",
            "825 resumenes articulo-entidad.",
        ],
    ),
    (
        "Frase para lucirte",
        [
            "La fortaleza del metodo no es que el algoritmo cuente mas rapido que una persona. La fortaleza es que cada conteo queda trazado a una evidencia textual auditable, lo que permite revisar, corregir y justificar decisiones metodologicas dentro de una revision sistematica.",
        ],
    ),
]


glossary = [
    ("Corpus", "Conjunto completo de articulos analizados. En esta presentacion son 79 articulos."),
    ("Mencion", "Aparicion textual de una palabra, sinonimo o variante de un contaminante o enfermedad."),
    ("Entidad", "Concepto normalizado detectado: por ejemplo, 'cadmio', 'aluminio', 'Alzheimer' o 'Parkinson'."),
    ("Seccion", "Parte del articulo donde aparece la mencion: titulo, resumen, introduccion, metodos, resultados, discusion, conclusion o referencias."),
    ("Ventana textual", "Fragmento alrededor de la mencion que se conserva para auditoria humana."),
    ("Pistas linguisticas", "Palabras cercanas que sugieren exposicion, dosis, medicion, modelo, asociacion, especulacion o negacion."),
    ("Confianza", "Nivel operativo asignado por reglas. Alta cuando hay evidencia directa en secciones fuertes; media cuando hay contexto asociativo; baja cuando es ambigua o bibliografica."),
    ("Asociacion", "Relacion textual contaminante-enfermedad con evidencia. Puede ser fuerte, debil, especulativa o insuficiente."),
    ("Peso de evidencia", "Score operativo calculado como fuerza de asociacion por confianza. No es una escala clinica validada universalmente."),
    ("K-Means", "Algoritmo no supervisado que agrupa articulos por similitud de perfiles. Sirve para exploracion y triage, no para probar causalidad."),
]


slides = [
    {
        "n": 1,
        "title": "Portada y escala del analisis",
        "claim": "La diapositiva abre con el alcance del estudio: 79 articulos, 14,739 menciones y 936 relaciones con evidencia.",
        "what": [
            "Presenta el tema general: contaminantes en agua y enfermedades neurodegenerativas.",
            "Anuncia que no se trata solo de una revision narrativa, sino de una analitica visual basada en un protocolo de deteccion y contextualizacion de palabras.",
            "Los tres numeros inferiores funcionan como prueba de escala: corpus, volumen de menciones y relaciones extraidas con evidencia textual.",
        ],
        "math": [
            "79 articulos = numero de documentos procesados por el pipeline.",
            "14,739 menciones = suma de todas las menciones detectadas de contaminantes y enfermedades.",
            "936 relaciones = pares contaminante-enfermedad generados solo cuando el sistema encontro evidencia textual cercana.",
        ],
        "say": "Esta primera diapositiva establece el tamano del corpus y la logica del trabajo. No estamos mostrando opiniones ni asociaciones inventadas, sino una estructura de evidencia rastreable desde cada grafica hasta fragmentos textuales de los articulos.",
        "importance": [
            "Da credibilidad porque muestra escala y trazabilidad.",
            "Prepara a la audiencia para entender que las figuras son productos de un pipeline reproducible.",
            "Permite aclarar desde el inicio que el objetivo es priorizar revision manual y no declarar causalidad.",
        ],
        "caution": [
            "No digas que 936 relaciones equivalen a 936 pruebas causales. Son relaciones textuales con evidencia, clasificadas por nivel de confianza.",
        ],
    },
    {
        "n": 2,
        "title": "Algoritmo de conteo de palabras con contexto",
        "claim": "El pipeline transforma PDFs o textos cientificos en salidas auditables mediante tres bloques: preparacion, deteccion y conteo, y salidas estructuradas.",
        "what": [
            "La parte A prepara el texto: entrada, extraccion, normalizacion y segmentacion por secciones.",
            "La parte B detecta entidades con lexicos de contaminantes y enfermedades, despues agrega contexto.",
            "La parte C genera clasificaciones y salidas auditables: rol del contaminante, confianza, asociacion, evidencia textual y tablas/graficos.",
        ],
        "math": [
            "El conteo inicial es una agrupacion por entidad: count(entidad).",
            "El contexto agrega variables adicionales: seccion, ventana textual, pistas linguisticas, cercania entre contaminante y enfermedad y negacion.",
            "La decision contextual puede verse como una funcion: decision = f(seccion, ventana, pistas, cercania, negacion).",
        ],
        "say": "Aqui esta el corazon metodologico. Primero se cuenta de manera objetiva, pero despues el algoritmo pregunta donde aparece la mencion, que palabras la rodean y si existe lenguaje de exposicion, medicion o asociacion. Por eso no tratamos igual una palabra en referencias que una palabra en metodos o resultados.",
        "importance": [
            "Muestra que el algoritmo es auditable y reproducible.",
            "Evita que la audiencia piense que solo se hizo un Ctrl+F masivo.",
            "Explica en que momento aparece el entendimiento de contexto: despues del match lexical, cuando se analiza seccion, ventana y pistas.",
        ],
        "caution": [
            "El algoritmo no entiende como un humano completo. Usa reglas explicitas de contexto. Esa es una ventaja para auditoria, porque se puede revisar por que tomo cada decision.",
        ],
    },
    {
        "n": 3,
        "title": "De mencion cruda a decision contextual",
        "claim": "Una palabra detectada se convierte en una clasificacion solo despues de revisar seccion, ventana textual, pistas y negacion.",
        "what": [
            "La diapositiva hace zoom al bloque contextual: una mencion como 'aluminio' entra al sistema y se evalua.",
            "Si aparece en metodos o resultados con pistas de exposicion, dosis o medicion, puede elevarse a exposicion principal.",
            "Si aparece solo en referencias, introduccion o contexto especulativo, se conserva como mencion, pero no como relacion fuerte.",
        ],
        "math": [
            "La clasificacion se modela como una regla compuesta: clasificacion = f(seccion, ventana_textual, pistas_exposicion, pistas_asociacion, cercania_entidades, negacion).",
            "La confianza alta requiere coincidencia entre seccion fuerte y evidencia directa.",
            "La confianza baja aparece cuando el texto es ambiguo, bibliografico, especulativo o con negacion.",
        ],
        "say": "Esta diapositiva responde a la pregunta mas importante: como evita el algoritmo falsos positivos. La palabra por si sola no decide. El sistema conserva la oracion, revisa si la mencion esta en una seccion fuerte y busca pistas como exposure, dose, measured, risk, associated o not.",
        "importance": [
            "Es la diapositiva que mas te conviene dominar si te preguntan por metodologia.",
            "Defiende la separacion entre extraccion literal e inferencia contextual.",
            "Justifica por que el sistema puede apoyar una revision sistematica: cada decision conserva evidencia textual.",
        ],
        "caution": [
            "No lo presentes como inteligencia artificial generativa que interpreta libremente. Es un sistema de reglas y NLP auditable.",
            "Si el fragmento no respalda la relacion, la salida debe quedar como evidencia insuficiente.",
        ],
    },
    {
        "n": 4,
        "title": "Menciones totales por contaminante",
        "claim": "La grafica muestra cuantas veces aparece cada contaminante en el corpus, no cuantos articulos estudian ese contaminante.",
        "what": [
            "Cada barra representa apariciones textuales detectadas para un contaminante.",
            "Los contaminantes con mas menciones fueron cadmio, aluminio, BMAA, pesticidas/plaguicidas, arsenico, mercurio, cobre y plomo.",
            "Esta grafica mide volumen textual, no relevancia experimental ni causalidad.",
        ],
        "math": [
            "Para cada contaminante i: menciones_i = numero de filas en mentions.csv donde entity_type = contaminant y label_es = i.",
            "El porcentaje de cada barra se calcula sobre el total de menciones de contaminantes, que es 9,831.",
            "Ejemplos del analisis actual: cadmio 1,383 menciones; aluminio 1,237; BMAA 753; pesticidas/plaguicidas 567; arsenico 506.",
        ],
        "say": "Este grafico es una capa basal. Si aluminio aparece muchas veces, significa que el termino fue mencionado muchas veces en el corpus; no significa automaticamente que 1,237 articulos lo estudiaron ni que todos los articulos lo trabajaron experimentalmente.",
        "importance": [
            "Ayuda a detectar que contaminantes dominan el lenguaje del corpus.",
            "Sirve para priorizar inspeccion, pero debe cruzarse con seccion y evidencia textual.",
            "Es util para explicar por que el conteo crudo es necesario, pero insuficiente.",
        ],
        "caution": [
            "Un articulo de revision largo puede inflar menciones.",
            "Las referencias tambien pueden inflar el conteo.",
            "No confundas menciones con numero de articulos.",
        ],
    },
    {
        "n": 5,
        "title": "Menciones totales de contaminantes por seccion",
        "claim": "La seccion revela donde se concentra el lenguaje sobre contaminantes y permite distinguir evidencia central de menciones perifericas.",
        "what": [
            "La grafica suma menciones de contaminantes por seccion del articulo.",
            "En este analisis, resultados concentra 3,165 menciones, referencias 2,328, introduccion 1,372, discusion 1,124 y metodos 890.",
            "Titulo, resumen y conclusion aportan menos menciones absolutas, pero pueden ser muy informativos por su posicion en el articulo.",
        ],
        "math": [
            "Para cada seccion s: menciones_s = suma de menciones de contaminantes detectadas en esa seccion.",
            "Resultados representa 32.2% de las menciones de contaminantes; referencias 23.7%; introduccion 14.0%; discusion 11.4%; metodos 9.1%.",
            "Esta distribucion ayuda a ponderar la evidencia, porque metodos/resultados suelen ser secciones mas fuertes que introduccion/referencias.",
        ],
        "say": "La seccion es el primer filtro de calidad. Si un contaminante aparece mucho en resultados o metodos, es mas probable que este relacionado con el analisis del estudio. Si aparece sobre todo en referencias, probablemente sea una mencion bibliografica y requiere mas cautela.",
        "importance": [
            "Conecta el conteo de palabras con interpretacion metodologica.",
            "Permite explicar por que no todas las menciones tienen el mismo valor.",
            "Sirve como control contra falsos positivos.",
        ],
        "caution": [
            "Una mencion en resultados no prueba causalidad; solo aumenta la probabilidad de que sea parte del analisis del articulo.",
            "Algunas revisiones tienen muchas referencias y pueden concentrar terminos alli.",
        ],
    },
    {
        "n": 6,
        "title": "Distribucion por seccion de los 25 contaminantes con mas menciones",
        "claim": "El heatmap muestra el perfil seccional de cada contaminante y ayuda a detectar si su presencia es experimental, narrativa o bibliografica.",
        "what": [
            "Las filas son contaminantes y las columnas son secciones del articulo.",
            "Cada celda indica cuantas menciones de ese contaminante aparecieron en esa seccion.",
            "Los colores mas intensos muestran concentraciones mayores de menciones.",
        ],
        "math": [
            "La celda (i, s) = count(contaminante i en seccion s).",
            "Se seleccionan los 25 contaminantes con mayor numero total de menciones para mantener legibilidad.",
            "La suma por fila recupera el total de menciones del contaminante; la suma por columna recupera el total de menciones en la seccion.",
        ],
        "say": "Esta figura mejora la lectura del grafico anterior porque ya no solo sabemos que seccion tiene mas menciones, sino que contaminante esta generando ese patron. Por ejemplo, un contaminante concentrado en referencias debe auditarse con mas cuidado que uno concentrado en metodos/resultados.",
        "importance": [
            "Permite identificar contaminantes con presencia robusta en secciones fuertes.",
            "Ayuda a planear auditoria manual por contaminante y seccion.",
            "Es una figura clave para explicar que el algoritmo no trata el texto como una bolsa de palabras sin estructura.",
        ],
        "caution": [
            "Un color fuerte indica frecuencia, no calidad de evidencia.",
            "Debe interpretarse junto con rol del contaminante, evidencia textual y tipo de estudio.",
        ],
    },
    {
        "n": 7,
        "title": "Numero de articulos donde el contaminante fue relevante",
        "claim": "Esta grafica cambia la unidad de analisis: pasa de menciones totales a articulos unicos.",
        "what": [
            "Cada barra indica en cuantos articulos aparece un contaminante como entidad relevante.",
            "Los contaminantes con mayor presencia por articulo fueron pesticidas/plaguicidas en 40 articulos, cobre en 33, cadmio en 32, plomo en 30, aluminio en 29, metales pesados en 29, mercurio en 28 y solventes en 28.",
            "Esta lectura es mas cercana a cobertura documental que a volumen textual.",
        ],
        "math": [
            "Para cada contaminante i: articulos_i = numero de article_id unicos donde i aparece en entity_summaries.csv.",
            "A diferencia de la diapositiva 4, un articulo cuenta una sola vez aunque mencione el contaminante muchas veces.",
            "Esto reduce el sesgo producido por articulos largos o revisiones extensas.",
        ],
        "say": "Aqui ya no preguntamos cuantas veces aparece una palabra, sino en cuantos articulos aparece. Por eso esta figura responde mejor a una pregunta de revision sistematica: que contaminantes estan distribuidos en mas documentos del corpus.",
        "importance": [
            "Ayuda a priorizar contaminantes con cobertura amplia.",
            "Complementa la frecuencia cruda porque reduce la influencia de documentos largos.",
            "Es util para decidir que familias de contaminantes revisar primero.",
        ],
        "caution": [
            "Que un contaminante aparezca en muchos articulos no significa que todos lo usen como exposicion principal.",
            "Debe cruzarse con el rol: exposicion principal, variable secundaria, revision o solo mencion.",
        ],
    },
    {
        "n": 8,
        "title": "Frecuencia de enfermedades, nivel de asociacion y tipo de estudio",
        "claim": "La diapositiva resume tres dimensiones del corpus: enfermedades detectadas, fuerza de asociacion y contexto de estudio.",
        "what": [
            "El primer grafico muestra la frecuencia documental de enfermedades neurodegenerativas.",
            "El segundo grafico separa las relaciones por nivel de asociacion: fuerte, debil, especulativa o sin evidencia suficiente.",
            "El tercer grafico muestra el tipo de estudio o contexto donde aparece la evidencia.",
        ],
        "math": [
            "En frecuencia por articulos, neurodegeneracion general aparece en 66 articulos, Alzheimer en 61, Parkinson en 59, demencia en 42, esclerosis lateral amiotrofica en 29 y deterioro cognitivo en 23.",
            "En relaciones, el corpus contiene 350 relaciones clasificadas como sin evidencia suficiente, 331 como asociacion debil, 149 como asociacion fuerte y 106 como mencion especulativa.",
            "Por tipo de articulo, el corpus incluye 31 experimentales in vitro, 21 epidemiologicos humanos, 14 revisiones/sintesis, 9 no determinados y 4 experimentales in vivo.",
        ],
        "say": "Esta diapositiva le da contexto a toda la analitica. No solo importa que contaminantes aparecen, sino con que enfermedades se vinculan, con que nivel de evidencia textual y en que tipo de estudio se reportan.",
        "importance": [
            "Ayuda a mostrar la heterogeneidad del corpus.",
            "Evita sobreinterpretar una matriz de asociaciones sin saber si predominan menciones fuertes o debiles.",
            "Sirve para identificar donde se requiere auditoria manual mas estricta.",
        ],
        "caution": [
            "Si predominan relaciones debiles o insuficientes, la conclusion debe ser prudente.",
            "El tipo de estudio afecta la interpretacion: una revision, un estudio epidemiologico y un ensayo in vitro no tienen el mismo peso metodologico.",
        ],
    },
    {
        "n": 9,
        "title": "Matriz basal contaminante-enfermedad",
        "claim": "La matriz convierte relaciones articulo-entidad en un mapa de calor de pares contaminante-enfermedad.",
        "what": [
            "Las filas agrupan categorias de contaminantes y las columnas agrupan enfermedades o desenlaces neurodegenerativos.",
            "Cada celda indica cuantas relaciones se detectaron con soporte textual.",
            "Las celdas mas intensas son pares con mayor densidad de relaciones en el corpus.",
        ],
        "math": [
            "celda(i, j) = numero de relaciones donde contaminant_category = i y disease_category = j.",
            "Ejemplo del analisis: metales pesados concentran 143 relaciones con Alzheimer, 123 con neurodegeneracion general, 120 con Parkinson y 79 con demencia.",
            "Pesticidas acumulan 50 relaciones con neurodegeneracion general y 46 con Parkinson.",
        ],
        "say": "Esta matriz es un mapa de densidad, no una prueba causal. Una celda alta significa que ese par aparece repetidamente con evidencia textual en el corpus. La interpretacion fuerte viene despues, cuando revisamos seccion, confianza y fragmentos de evidencia.",
        "importance": [
            "Resume muchas relaciones en una estructura facil de leer.",
            "Permite identificar rapidamente que familias de contaminantes dominan ciertos desenlaces.",
            "Sirve como punto de partida para seleccionar pares prioritarios.",
        ],
        "caution": [
            "La matriz cuenta relaciones, no calidad metodologica.",
            "Una celda alta puede reflejar articulos de revision o lenguaje general.",
        ],
    },
    {
        "n": 10,
        "title": "Agregacion por categoria y peso de evidencia",
        "claim": "La matriz ponderada combina fuerza de asociacion y confianza para comparar familias de contaminantes.",
        "what": [
            "Agrupa contaminantes por categoria para reducir ruido terminologico.",
            "En lugar de contar todas las relaciones por igual, suma un peso operativo.",
            "El resultado destaca categorias con mayor acumulacion de evidencia ponderada.",
        ],
        "math": [
            "peso = fuerza x confianza.",
            "Fuerza: fuerte = 3, debil = 2, especulativa = 1, insuficiente = 0.",
            "Confianza: alta = 3, media = 2, baja = 1.",
            "Ejemplos del analisis ponderado: metales pesados tienen peso 435 con Alzheimer, 426 con neurodegeneracion general y 318 con Parkinson. Pesticidas tienen peso 143 con Parkinson y 108 con neurodegeneracion general.",
        ],
        "say": "Esta diapositiva mejora la matriz basal porque no todas las relaciones valen igual. Una relacion fuerte y de alta confianza pesa mas que una mencion especulativa. Asi obtenemos una priorizacion mas util para decidir que pares auditar primero.",
        "importance": [
            "Reduce ruido por sinonimos y terminos individuales.",
            "Integra cantidad y calidad contextual en un mismo score.",
            "Es util para ordenar trabajo de extraccion y validacion manual.",
        ],
        "caution": [
            "Estos pesos son una escala operativa del pipeline, no una escala clinica universal validada.",
            "Conviene reportar los pesos como decision metodologica auditable y, si se publica, hacer analisis de sensibilidad.",
        ],
    },
    {
        "n": 11,
        "title": "Grafico de burbujas",
        "claim": "La burbuja separa volumen de evidencia y fuerza promedio para cada par contaminante-enfermedad.",
        "what": [
            "Cada burbuja representa un par contaminante-enfermedad.",
            "El area de la burbuja representa la suma de pesos.",
            "El color representa la fuerza promedio del par.",
        ],
        "math": [
            "area_burbuja proporcional a sum(peso).",
            "color proporcional a avg_strength = weighted_score / n_relations.",
            "Fila = contaminante; columna = enfermedad.",
        ],
        "say": "Esta figura ayuda a no confundir cantidad con fuerza. Una burbuja grande puede deberse a muchas relaciones moderadas, mientras que una burbuja mas pequena pero oscura puede indicar pocas relaciones con mayor fuerza promedio.",
        "importance": [
            "Permite priorizar pares grandes y oscuros para auditoria manual temprana.",
            "Ayuda a separar acumulacion textual de intensidad contextual.",
            "Es mas informativa que una barra simple porque muestra dos variables a la vez.",
        ],
        "caution": [
            "No todos los pares grandes son necesariamente los mas concluyentes.",
            "Debe revisarse el fragmento textual asociado antes de redactar conclusiones cientificas.",
        ],
    },
    {
        "n": 12,
        "title": "Red bipartita",
        "claim": "La red muestra como se conectan contaminantes y enfermedades mediante relaciones con evidencia textual.",
        "what": [
            "Los nodos del lado izquierdo son contaminantes y los del lado derecho son enfermedades.",
            "Un enlace existe cuando el pipeline encontro una relacion contaminante-enfermedad con evidencia.",
            "El grosor del enlace representa mayor peso de evidencia.",
        ],
        "math": [
            "grado(nodo) = numero de enlaces del nodo.",
            "Los contaminantes con mayor grado en la salida completa incluyen cadmio, mercurio, plomo, cobre, zinc y manganeso, cada uno conectado con ocho categorias de desenlace.",
            "Por score total, aluminio, cadmio, mercurio, pesticidas/plaguicidas y plomo son nodos relevantes en la red.",
        ],
        "say": "La red es una forma de ver estructura. No solo nos dice que pares son frecuentes, sino que contaminantes funcionan como conectores transversales a multiples desenlaces neurodegenerativos.",
        "importance": [
            "Ayuda a detectar contaminantes que aparecen en muchas relaciones diferentes.",
            "Permite distinguir nodos especificos de nodos muy generales.",
            "Es util para decidir si una subseccion de la revision debe organizarse por contaminante o por enfermedad.",
        ],
        "caution": [
            "Una red densa puede reflejar terminos generales, revisiones amplias o literatura muy repetida.",
            "El grado alto no implica mayor causalidad.",
        ],
    },
    {
        "n": 13,
        "title": "Contexto por seccion como filtro contra falsos positivos",
        "claim": "La seccion del articulo es una variable de calidad para interpretar asociaciones.",
        "what": [
            "El heatmap cruza nivel de asociacion con seccion del articulo.",
            "Muestra donde se sostienen las asociaciones fuertes, debiles, especulativas o insuficientes.",
            "La figura justifica por que el algoritmo da mas peso a metodos/resultados que a introduccion/referencias.",
        ],
        "math": [
            "celda(a, s) = numero de relaciones con association = a y section = s.",
            "En el analisis, asociacion fuerte se concentra en resultados, titulo, resumen y metodos: resultados 79, resumen 33, titulo 22 y metodos 15.",
            "Asociacion debil e insuficiente aparecen frecuentemente en referencias e introduccion, lo que exige cautela.",
        ],
        "say": "Esta figura es central para defender el control de falsos positivos. Si una relacion aparece en metodos o resultados, tiene una lectura distinta a una relacion que aparece solo en referencias. Por eso el pipeline no solo detecta entidades; tambien conserva la seccion.",
        "importance": [
            "Muestra donde esta apoyada la evidencia.",
            "Permite auditar si una asociacion depende de secciones fuertes o perifericas.",
            "Ayuda a separar exposicion principal, variable secundaria y simple mencion bibliografica.",
        ],
        "caution": [
            "No se debe descartar automaticamente la introduccion o discusion, pero si se debe reducir la confianza cuando no hay evidencia directa.",
            "El peso por seccion es una regla metodologica del pipeline y debe declararse como tal.",
        ],
    },
    {
        "n": 14,
        "title": "Pares prioritarios",
        "claim": "La clasificacion ordena los pares contaminante-enfermedad que conviene auditar primero.",
        "what": [
            "La grafica presenta pares con mayor score ponderado.",
            "Los pares principales incluyen aluminio-Alzheimer, aluminio-demencia, manganeso-neurodegeneracion general, cobre-neurodegeneracion general, pesticidas/plaguicidas-neurodegeneracion general, paraquat-Parkinson y pesticidas/plaguicidas-Parkinson.",
            "Cada barra resume peso, numero de relaciones y fuerza promedio.",
        ],
        "math": [
            "score(par) = suma de fuerza x confianza en todas las relaciones del par.",
            "n_relations = numero de repeticiones detectadas.",
            "avg_strength = score / n_relations.",
            "Ejemplos: aluminio-Alzheimer tiene score 108 con 17 relaciones; aluminio-demencia score 88 con 16 relaciones; paraquat-Parkinson score 51 con 9 relaciones.",
        ],
        "say": "Esta figura convierte el analisis en una agenda de revision. Los pares de arriba no son conclusiones finales, sino candidatos para auditoria manual: revisar articulos, secciones y fragmentos de evidencia antes de redactar inferencias cientificas.",
        "importance": [
            "Ayuda a organizar el trabajo del equipo.",
            "Prioriza pares donde hay acumulacion de evidencia textual y confianza.",
            "Facilita pasar de visualizacion a decisiones practicas de revision sistematica.",
        ],
        "caution": [
            "Un score alto puede venir de muchas relaciones moderadas o de pocas relaciones fuertes.",
            "Siempre revisar score y n_relations juntos para evitar sobreinterpretar outliers.",
        ],
    },
    {
        "n": 15,
        "title": "PCA + K-Means exploratorio, y nota de actualizacion a K-Means",
        "claim": "La diapositiva original muestra clustering exploratorio de articulos; la version actual del pipeline ya fue ajustada para usar K-Means sin PCA.",
        "what": [
            "La diapositiva explica que cada articulo se representa como un vector de rasgos: contaminantes, enfermedades, secciones, roles y tipos de asociacion.",
            "En la version original, PCA proyectaba esa matriz a dos dimensiones y K-Means agrupaba articulos por similitud.",
            "En la version actual del pipeline, se elimino PCA como parte activa del algoritmo y se usa K-Means sobre una matriz estandarizada de rasgos compactos e interpretables.",
        ],
        "math": [
            "Matriz original: X = articulo x rasgo.",
            "Transformacion: X' = log(1 + X) para reducir la dominancia de conteos muy grandes.",
            "K-Means busca k centroides y asigna cada articulo al centroide mas cercano.",
            "En la salida actual, k = 4 clusters: cluster 0 con 53 articulos, cluster 1 con 7, cluster 2 con 3 y cluster 3 con 16.",
        ],
        "say": "Esta figura no busca demostrar relaciones causales. Su funcion es agrupar articulos por perfiles similares para hacer triage: articulos con patrones parecidos de contaminantes, enfermedades, secciones y asociaciones pueden revisarse juntos.",
        "importance": [
            "Ayuda a detectar articulos atipicos o grupos tematicos.",
            "Sirve para organizar auditoria manual, no para validar una hipotesis biologica.",
            "Permite explicar por que K-Means es util: reduce un conjunto complejo de rasgos a grupos exploratorios revisables.",
        ],
        "caution": [
            "Si presentas esta diapositiva tal como esta, aclara que PCA fue una visualizacion exploratoria previa.",
            "Como el pipeline actual ya esta en K-Means-only, la frase mas segura es: 'La version actual agrupa con K-Means usando rasgos interpretables; PCA ya no se usa para tomar decisiones'.",
            "No digas que los clusters prueban asociaciones contaminante-enfermedad. Agrupan articulos, no pares causales.",
        ],
    },
    {
        "n": 16,
        "title": "Recomendacion metodologica",
        "claim": "La estrategia mas defendible es reglas auditables ahora y aprendizaje supervisado despues, cuando existan etiquetas humanas.",
        "what": [
            "La diapositiva sintetiza cinco componentes: lexicos, contexto, reglas, matrices y machine learning.",
            "Defiende una aproximacion hibrida: diccionarios y reglas para trazabilidad; visual analytics para priorizar; ML para una fase futura.",
            "Cierra con una conclusion prudente: usar figuras para priorizar auditoria manual, no para cerrar conclusiones sin respaldo textual.",
        ],
        "math": [
            "La parte de reglas produce variables discretas: seccion, rol, confianza, asociacion.",
            "Las matrices y scores permiten sumar evidencia con reglas explicitas.",
            "El ML futuro requeriria un conjunto etiquetado por humanos: exposicion principal, variable secundaria, mencion bibliografica, asociacion fuerte/debil, etc.",
        ],
        "say": "La razon para empezar con reglas y NLP es la auditoria. En una revision sistematica necesitamos justificar cada decision. Mas adelante, cuando tengamos etiquetas humanas, podemos entrenar un modelo supervisado para aprender patrones mas complejos, pero sin perder trazabilidad.",
        "importance": [
            "Deja claro que el metodo es gradual y cientificamente prudente.",
            "Evita vender K-Means o PCA como solucion final.",
            "Conecta la analitica visual con el flujo real de una revision sistematica.",
        ],
        "caution": [
            "El siguiente paso ideal no es meter mas modelos sin control, sino validar manualmente una muestra y medir precision, recall y acuerdo entre revisores.",
        ],
    },
]


hard_questions = [
    (
        "¿Esto prueba que un contaminante causa una enfermedad?",
        "No. El pipeline detecta y organiza evidencia textual. Puede identificar asociaciones reportadas o sugeridas, pero la causalidad requiere evaluacion metodologica, diseno del estudio, control de confusores y revision critica.",
    ),
    (
        "¿Por que no usar directamente un modelo de inteligencia artificial generativa?",
        "Porque en una revision sistematica la trazabilidad es prioritaria. Las reglas permiten auditar cada decision. Un modelo generativo podria apoyar despues, pero no debe reemplazar la evidencia textual verificable.",
    ),
    (
        "¿Los pesos de seccion y confianza estan publicados tal cual?",
        "No como una escala universal. Son una decision metodologica operativa, inspirada en principios de jerarquia de evidencia textual: metodos/resultados pesan mas que introduccion/referencias. Por eso deben reportarse y validarse con sensibilidad.",
    ),
    (
        "¿Por que K-Means?",
        "Porque agrupa articulos sin necesidad de etiquetas previas y ayuda a detectar perfiles documentales. Sirve para triage exploratorio, no para clasificacion final ni causalidad.",
    ),
    (
        "¿Que haria falta para usar aprendizaje supervisado?",
        "Un conjunto de articulos o fragmentos etiquetados manualmente por revisores: contaminante usado, rol del contaminante, tipo de estudio, enfermedad asociada y fuerza de asociacion. Con eso se podria entrenar y evaluar un clasificador.",
    ),
    (
        "¿Como se controlan falsos positivos?",
        "Con seccion, ventana textual, pistas linguisticas, negacion, evidencia textual y auditoria manual. Una palabra en referencias no se trata igual que una exposicion descrita en metodos/resultados.",
    ),
]


closing_script = [
    "Como cierre, yo enfatizaria que esta analitica no sustituye al criterio experto. Lo que hace es ordenar el corpus y hacer transparente donde estan las menciones, donde esta la evidencia textual y que pares contaminante-enfermedad merecen revision prioritaria.",
    "La ruta mas solida es mantener el pipeline reproducible, auditar manualmente los pares prioritarios y despues usar esas etiquetas para mejorar el sistema con aprendizaje supervisado.",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_label_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.6)
        cells[1].width = Inches(4.9)
        for c in cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
        set_cell_shading(cells[0], "E8EEF5")
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value)
    doc.add_paragraph()


def configure_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Title", 20, "1F4D78"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def build_docx():
    GUIDES_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    r.italic = True
    r.font.size = Pt(12)

    add_para(doc, "Fuente analizada: analitica_contaminantes.pptx")
    add_para(doc, "Uso sugerido: guia de estudio para exponer ante investigadores del proyecto.")
    add_para(doc, "Idea rectora: conteo reproducible + contexto textual + evidencia auditable + visualizacion para priorizar revision manual.")

    doc.add_heading("Como usar esta guia", level=1)
    add_bullets(
        doc,
        [
            "Lee primero la tesis central y memoriza el mensaje de 30 segundos.",
            "Despues estudia las diapositivas 2 y 3, porque ahi esta la defensa metodologica del algoritmo.",
            "Para las diapositivas 4 a 14, practica explicar que mide cada grafica y que no mide.",
            "En la diapositiva 15, usa la nota de actualizacion: el pipeline actual ya usa K-Means sin PCA.",
            "Cierra con la diapositiva 16: reglas auditables ahora, aprendizaje supervisado despues.",
        ],
    )

    for heading, paragraphs in overview_sections:
        doc.add_heading(heading, level=1)
        for para in paragraphs:
            add_para(doc, para)

    doc.add_heading("Glosario rapido", level=1)
    add_label_table(doc, glossary)

    doc.add_heading("Guia diapositiva por diapositiva", level=1)
    for slide in slides:
        doc.add_page_break()
        doc.add_heading(f"Diapositiva {slide['n']}. {slide['title']}", level=1)
        img = SLIDE_DIR / f"source-slide-{slide['n']:02d}.png"
        if img.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img), width=Inches(6.4))
        add_label_table(
            doc,
            [
                ("Idea central", slide["claim"]),
                ("Guion breve", slide["say"]),
            ],
        )
        doc.add_heading("Que muestra", level=2)
        add_bullets(doc, slide["what"])
        doc.add_heading("Lectura matematica o estadistica", level=2)
        add_bullets(doc, slide["math"])
        doc.add_heading("Por que importa para la revision", level=2)
        add_bullets(doc, slide["importance"])
        doc.add_heading("Cuidado al presentarla", level=2)
        add_bullets(doc, slide["caution"])

    doc.add_page_break()
    doc.add_heading("Preguntas dificiles y respuestas elegantes", level=1)
    for question, answer in hard_questions:
        doc.add_heading(question, level=2)
        add_para(doc, answer)

    doc.add_heading("Cierre sugerido", level=1)
    for para in closing_script:
        add_para(doc, para)

    doc.add_heading("Frases de alto impacto para la exposicion", level=1)
    add_bullets(
        doc,
        [
            "Conteo no es evidencia; evidencia es conteo contextualizado y auditable.",
            "La seccion del articulo funciona como filtro de calidad textual.",
            "Las figuras no sustituyen la lectura critica; ordenan donde debe empezar la lectura critica.",
            "K-Means agrupa articulos por perfil de evidencia, no prueba causalidad.",
            "El siguiente salto metodologico es entrenar modelos supervisados con etiquetas humanas, no abandonar la trazabilidad.",
        ],
    )

    doc.save(DOCX_OUT)


def md_bullets(items):
    return "\n".join(f"- {item}" for item in items)


def build_markdown():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    lines = [f"# {TITLE}", "", SUBTITLE, "", "**Fuente:** `analitica_contaminantes.pptx`", ""]
    lines.append("## Como usar esta guia")
    lines.append(md_bullets([
        "Memoriza primero la tesis central.",
        "Domina las diapositivas 2 y 3 para defender el algoritmo.",
        "En cada grafica, explica que mide y que no mide.",
        "En la diapositiva 15, aclara que el pipeline actual ya usa K-Means sin PCA.",
    ]))
    lines.append("")
    for heading, paragraphs in overview_sections:
        lines.append(f"## {heading}")
        lines.extend(paragraphs)
        lines.append("")
    lines.append("## Glosario rapido")
    lines.append("| Termino | Explicacion |")
    lines.append("|---|---|")
    for term, explanation in glossary:
        lines.append(f"| {term} | {explanation} |")
    lines.append("")
    lines.append("## Guia diapositiva por diapositiva")
    for slide in slides:
        img = SLIDE_DIR / f"source-slide-{slide['n']:02d}.png"
        lines.append(f"### Diapositiva {slide['n']}. {slide['title']}")
        if img.exists():
            lines.append(f"![Diapositiva {slide['n']}]({img})")
        lines.append(f"**Idea central:** {slide['claim']}")
        lines.append("")
        lines.append(f"**Guion breve:** {slide['say']}")
        lines.append("")
        lines.append("**Que muestra**")
        lines.append(md_bullets(slide["what"]))
        lines.append("")
        lines.append("**Lectura matematica o estadistica**")
        lines.append(md_bullets(slide["math"]))
        lines.append("")
        lines.append("**Por que importa para la revision**")
        lines.append(md_bullets(slide["importance"]))
        lines.append("")
        lines.append("**Cuidado al presentarla**")
        lines.append(md_bullets(slide["caution"]))
        lines.append("")
    lines.append("## Preguntas dificiles y respuestas elegantes")
    for question, answer in hard_questions:
        lines.append(f"### {question}")
        lines.append(answer)
        lines.append("")
    lines.append("## Cierre sugerido")
    lines.extend(closing_script)
    lines.append("")
    lines.append("## Frases de alto impacto")
    lines.append(md_bullets([
        "Conteo no es evidencia; evidencia es conteo contextualizado y auditable.",
        "La seccion del articulo funciona como filtro de calidad textual.",
        "Las figuras no sustituyen la lectura critica; ordenan donde debe empezar la lectura critica.",
        "K-Means agrupa articulos por perfil de evidencia, no prueba causalidad.",
        "El siguiente salto metodologico es entrenar modelos supervisados con etiquetas humanas, no abandonar la trazabilidad.",
    ]))
    MD_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    build_markdown()
    build_docx()
    print(MD_OUT)
    print(DOCX_OUT)
