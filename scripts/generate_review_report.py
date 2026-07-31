#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import unicodedata
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches


RELATION_COLUMNS = [
    "entity_a_label",
    "entity_b_label",
    "association",
    "confidence",
    "section",
    "evidence_text",
]


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def slugify(value: str) -> str:
    """Mismo slug ASCII que usan visualize.py y visual_analytics.py para nombrar figuras."""
    cleaned = unicodedata.normalize("NFD", value)
    cleaned = "".join(ch for ch in cleaned if unicodedata.category(ch) != "Mn")
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "_", cleaned).strip("_").lower()
    return cleaned or "variable"


def unslugify(slug: str) -> str:
    return slug.replace("_", " ").strip() or "la variable"


def read_protocol(output_dir: Path) -> dict[str, str]:
    """Nombres legibles del protocolo, escritos por export.py en review_miner_results.json."""
    protocol = {"name": "Protocolo sin nombre", "variable_a": "Variable A", "variable_b": "Variable B"}
    path = output_dir / "review_miner_results.json"
    if not path.exists():
        return protocol
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return protocol
    block = payload.get("protocol") or {}
    for key in protocol:
        value = str(block.get(key, "")).strip()
        if value:
            protocol[key] = value
    return protocol


def figure_descriptions(name_a: str, name_b: str) -> dict[str, str]:
    """Descripciones por figura, derivadas de los nombres de variables del protocolo."""
    lower_a = name_a.lower()
    lower_b = name_b.lower()
    return {
        f"frecuencia_{slugify(name_a)}": f"Frecuencia de menciones de {lower_a} en el corpus.",
        f"frecuencia_{slugify(name_b)}": f"Frecuencia de menciones de {lower_b} en el corpus.",
        "heatmap_asociaciones": f"Matriz de categorias {lower_a} × {lower_b} basada en relaciones textuales.",
        "tipo_estudio": "Distribucion de articulos por tipo de estudio inferido.",
        "nivel_asociacion": "Relaciones agrupadas por nivel de asociacion textual.",
        "association_network": f"Red exploratoria de relaciones {lower_a}-{lower_b}.",
        f"bubble_{slugify(name_a)}_{slugify(name_b)}": (
            f"Pares {lower_a}-{lower_b} representados como burbujas por peso/frecuencia."
        ),
        "top_association_pairs": "Pares de asociacion mas frecuentes.",
        "category_association_heatmap": f"Heatmap avanzado por categorias de {lower_a} y {lower_b}.",
        "association_by_section": "Distribucion de evidencia por seccion del articulo.",
        "cluster_sizes": "Tamanos de clusters en K-Means exploratorio.",
        "kmeans_cluster_map": "Mapa K-Means para triage exploratorio y deteccion de outliers.",
    }


def describe_figure(stem: str, descriptions: dict[str, str]) -> str:
    """Resuelve la descripcion por nombre exacto y, si no coincide, por prefijo de slot."""
    if stem in descriptions:
        return descriptions[stem]
    if stem.startswith("frecuencia_"):
        return f"Frecuencia de menciones de {unslugify(stem[len('frecuencia_'):])} en el corpus."
    if stem.startswith("bubble_"):
        return "Pares de entidades representados como burbujas por peso/frecuencia."
    return "Figura generada por el pipeline."


CELL_CHAR_LIMIT = 220


def _clip(value: str) -> str:
    """Acorta una celda marcando el corte.

    La evidencia textual suele superar el limite de la celda. Un corte mudo
    haria imposible saber que la cita continua, y es justo al final donde
    aparecerian una negacion o una matizacion.
    """

    if len(value) <= CELL_CHAR_LIMIT:
        return value
    return value[: CELL_CHAR_LIMIT - 4].rstrip() + " […]"


def add_metric_table(document: Document, metrics: dict[str, object]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Indicador"
    header[1].text = "Valor"
    for key, value in metrics.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)


def add_top_relations(document: Document, relations: pd.DataFrame, name_a: str, name_b: str) -> None:
    if relations.empty:
        document.add_paragraph(
            f"No se encontraron relaciones {name_a.lower()}-{name_b.lower()} con evidencia cercana."
        )
        return
    missing = [column for column in RELATION_COLUMNS if column not in relations.columns]
    if missing:
        raise KeyError(
            "relations.csv no contiene las columnas esperadas "
            f"{missing}; columnas presentes: {list(relations.columns)}"
        )
    headers = {
        "entity_a_label": name_a,
        "entity_b_label": name_b,
        "association": "Asociacion",
        "confidence": "Confianza",
        "section": "Seccion",
        "evidence_text": "Evidencia textual",
    }
    preview = relations[RELATION_COLUMNS].head(12)
    table = document.add_table(rows=1, cols=len(RELATION_COLUMNS))
    table.style = "Table Grid"
    for idx, column in enumerate(RELATION_COLUMNS):
        table.rows[0].cells[idx].text = headers[column]
    for _, data in preview.iterrows():
        row = table.add_row().cells
        for idx, column in enumerate(RELATION_COLUMNS):
            row[idx].text = _clip(str(data.get(column, "")))


def figure_rows(output_dir: Path, descriptions: dict[str, str]) -> list[tuple[str, str]]:
    figures = sorted(
        [
            path
            for path in output_dir.rglob("*")
            if path.suffix.lower() in {".svg", ".png", ".jpg", ".jpeg", ".webp"}
        ]
    )
    rows = []
    for figure in figures:
        stem = figure.stem
        rows.append((str(figure.relative_to(output_dir)), describe_figure(stem, descriptions)))
    return rows


def add_figure_references(document: Document, rows: list[tuple[str, str]]) -> None:
    if not rows:
        document.add_paragraph("No se encontraron figuras en la carpeta de salida.")
        return
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Figura"
    table.rows[0].cells[1].text = "Descripcion"
    for file_name, description in rows:
        row = table.add_row().cells
        row[0].text = file_name
        row[1].text = description


def build_report(output_dir: Path) -> Path:
    articles = read_csv(output_dir / "articles.csv")
    mentions = read_csv(output_dir / "mentions.csv")
    summaries = read_csv(output_dir / "entity_summaries.csv")
    relations = read_csv(output_dir / "relations.csv")

    protocol = read_protocol(output_dir)
    name_a = protocol["variable_a"]
    name_b = protocol["variable_b"]
    descriptions = figure_descriptions(name_a, name_b)
    figures = figure_rows(output_dir, descriptions)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    document.add_heading("Reporte de revision sistematica asistida", level=0)
    document.add_paragraph("Aplicacion: NEXO — Mineria de datos para revisiones.")
    document.add_paragraph(f"Protocolo: {protocol['name']}")
    document.add_paragraph(f"Variable A: {name_a} — Variable B: {name_b}")
    document.add_paragraph(f"Fecha de generacion: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph(f"Carpeta de salida: {output_dir}")

    document.add_heading("Resumen ejecutivo", level=1)
    metrics = {
        "Articulos procesados": len(articles),
        "Articulos con texto extraible": int((articles.get("text_extractable", pd.Series(dtype=str)).astype(str).str.lower() == "si").sum())
        if not articles.empty
        else 0,
        "Menciones auditables": len(mentions),
        "Resumenes articulo-entidad": len(summaries),
        "Relaciones con evidencia textual cercana": len(relations),
        "Figuras generadas": len(figures),
    }
    add_metric_table(document, metrics)

    document.add_heading("Metodo general", level=1)
    document.add_paragraph(
        f"El analisis aplica un pipeline auditable basado en lexicos controlados de {name_a.lower()} y "
        f"{name_b.lower()}, extraccion de texto desde PDFs, deteccion aproximada de secciones, ventanas de "
        "contexto, pistas de exposicion, dosis, asociacion, especulacion y negacion, ademas de relaciones "
        f"{name_a.lower()}-{name_b.lower()} por cercania textual."
    )
    document.add_paragraph(
        "Las salidas deben interpretarse como apoyo para revision sistematica y auditoria manual. Una mencion o "
        "co-ocurrencia no equivale a causalidad. Las asociaciones reportadas requieren confirmacion por lectura humana."
    )

    document.add_heading("Resultados principales", level=1)
    add_top_relations(document, relations, name_a, name_b)

    document.add_heading("Referencias a figuras generadas", level=1)
    add_figure_references(document, figures)

    document.add_heading("Advertencias metodologicas", level=1)
    document.add_paragraph(
        "Los resultados dependen de la calidad de extraccion de texto del PDF, de la cobertura de los lexicos y de la "
        "estructura textual del articulo. El sistema distingue mencion, co-ocurrencia y asociacion textual, pero no "
        "establece evidencia concluyente ni recomienda decisiones clinicas."
    )

    report_path = output_dir / "reporte_revision_nexo.docx"
    document.save(report_path)
    return report_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Genera reporte Word para las salidas de NEXO.")
    parser.add_argument("--output-dir", required=True, help="Carpeta de salida del pipeline.")
    args = parser.parse_args()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    report_path = build_report(output_dir)
    print(f"Reporte Word generado: {report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
